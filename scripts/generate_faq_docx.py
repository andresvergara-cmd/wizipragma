"""
Genera el documento FAQ mejorado para el Knowledge Base de Comfi.
Cada pregunta-respuesta está claramente separada para mejor indexación.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Título
title = doc.add_heading('Centro de Conocimiento - FAQ Comfama', level=0)

# Subtítulo
doc.add_paragraph('Documento de preguntas frecuentes para el asistente virtual Comfi de Comfama.')
doc.add_paragraph('---')

# FAQ Data - cada entrada es independiente para mejor chunking
faqs = [
    # AFILIACIÓN
    {
        "categoria": "AFILIACIÓN",
        "pregunta": "¿Cómo me puedo afiliar a Comfama?",
        "respuesta": "Puedes afiliarte a Comfama desde el portal de servicios en línea como trabajador independiente, pensionado o afiliando empleados domésticos. Debes completar el formulario de afiliación y adjuntar los documentos requeridos como cédula de ciudadanía y certificado laboral. La afiliación también puede ser realizada por tu empleador si eres trabajador dependiente del sector privado en Antioquia."
    },
    {
        "categoria": "AFILIACIÓN",
        "pregunta": "¿Cómo consultar el estado de mi afiliación a Comfama?",
        "respuesta": "Para consultar el estado de tu afiliación a Comfama, ingresa al portal de servicios en línea de Comfama (www.comfama.com) con tu documento de identidad y contraseña. En la sección de afiliación podrás verificar tu estado actual y tu categoría de afiliación (A, B o C). También puedes llamar a la línea de atención (604) 360 70 80."
    },
    {
        "categoria": "AFILIACIÓN",
        "pregunta": "¿Cómo consultar el estado de afiliación si soy beneficiario?",
        "respuesta": "Si eres beneficiario, puedes ingresar al portal de servicios de Comfama con el documento del afiliado principal y revisar el estado de afiliación. También puedes consultar llamando a la línea de atención (604) 360 70 80 o visitando una sede de Comfama."
    },
    {
        "categoria": "AFILIACIÓN",
        "pregunta": "¿A quiénes puedo afiliar como beneficiarios en Comfama?",
        "respuesta": "En Comfama puedes afiliar como beneficiarios a: tu cónyuge o compañero(a) permanente, hijos dependientes económicamente, padres que dependan económicamente de ti, y otros familiares según las condiciones establecidas por la normativa vigente. Cada beneficiario debe cumplir requisitos específicos de parentesco y dependencia económica."
    },
    {
        "categoria": "AFILIACIÓN",
        "pregunta": "¿Qué edad deben tener mis hijos para afiliarlos a Comfama?",
        "respuesta": "Generalmente los hijos pueden estar afiliados a Comfama hasta los 23 años, siempre que dependan económicamente del afiliado principal y cumplan con las condiciones educativas establecidas. Los hijos menores de 18 años se afilian automáticamente. Los hijos entre 18 y 23 años deben demostrar que estudian o dependen económicamente."
    },
    {
        "categoria": "AFILIACIÓN",
        "pregunta": "¿Puedo afiliarme a Comfama si soy pensionado?",
        "respuesta": "Sí, los pensionados pueden afiliarse voluntariamente a Comfama y acceder a servicios dependiendo del aporte realizado. Debes realizar la afiliación desde el portal de servicios en línea o en una sede de Comfama presentando tu documento de identidad y resolución de pensión."
    },
    {
        "categoria": "AFILIACIÓN",
        "pregunta": "¿Cómo afiliar mi grupo familiar a Comfama?",
        "respuesta": "Desde el portal de servicios de Comfama puedes registrar a cada beneficiario de tu grupo familiar cargando los documentos requeridos: documento de identidad del beneficiario, registro civil (para hijos), acta de matrimonio o declaración de unión marital (para cónyuge/compañero). También puedes hacerlo presencialmente en cualquier sede de Comfama."
    },
    {
        "categoria": "AFILIACIÓN",
        "pregunta": "¿Qué beneficios tiene la afiliación a Comfama?",
        "respuesta": "La afiliación a Comfama te permite acceder a múltiples beneficios: subsidios familiares en dinero, subsidio de vivienda, programas educativos y de formación, servicios de salud y bienestar, créditos con tasas preferenciales, recreación y turismo en centros vacacionales, programas culturales, bibliotecas, y programas sociales para toda tu familia."
    },
    {
        "categoria": "AFILIACIÓN",
        "pregunta": "¿Cómo afilio a mi empresa a Comfama?",
        "respuesta": "Para afiliar tu empresa a Comfama debes: 1) Ingresar al portal empresarial de Comfama (www.comfama.com/empresas) o acercarte a una sede. 2) Presentar el certificado de existencia y representación legal de la Cámara de Comercio, RUT actualizado, documento de identidad del representante legal, y formulario de afiliación empresarial diligenciado. 3) Una vez afiliada la empresa, debes reportar a tus trabajadores y realizar los aportes parafiscales mensuales del 4% sobre la nómina. También puedes llamar al (604) 360 70 80 para recibir asesoría personalizada sobre el proceso."
    },
    {
        "categoria": "AFILIACIÓN",
        "pregunta": "¿Cuál es la tarifa de afiliación a Comfama?",
        "respuesta": "La tarifa de afiliación a Comfama es el 4% del salario mensual del trabajador, que es aportado directamente por el empleador como parte de los aportes parafiscales. El trabajador no paga nada de su bolsillo. Ejemplo: si tu salario es de $2,000,000 COP, el aporte mensual es de $80,000 COP pagado por tu empleador. Para trabajadores independientes y pensionados que se afilian voluntariamente, las tarifas varían según el plan de servicios seleccionado."
    },
    {
        "categoria": "AFILIACIÓN",
        "pregunta": "¿Qué pasa si mi empleador no me afilia a Comfama?",
        "respuesta": "Si tu empleador no te afilia a una caja de compensación familiar como Comfama, está incumpliendo la ley colombiana. Todo empleador del sector privado está obligado a afiliar a sus trabajadores y pagar el aporte del 4% sobre la nómina. Puedes reportar esta situación ante el Ministerio del Trabajo o la Superintendencia del Subsidio Familiar. También puedes comunicarte con Comfama al (604) 360 70 80 para recibir orientación sobre cómo proceder."
    },
    # CERTIFICADOS Y DOCUMENTOS
    {
        "categoria": "CERTIFICADOS Y DOCUMENTOS",
        "pregunta": "¿Cómo descargar el certificado de afiliación de Comfama?",
        "respuesta": "Para descargar tu certificado de afiliación de Comfama, debes ingresar al portal de servicios en línea (www.comfama.com), ir a la sección de certificados y generar el certificado de afiliación. El documento se genera en formato PDF y puedes descargarlo e imprimirlo inmediatamente."
    },
    {
        "categoria": "CERTIFICADOS Y DOCUMENTOS",
        "pregunta": "¿Cómo descargar el carné de afiliación de Comfama?",
        "respuesta": "En el portal de servicios de Comfama puedes generar tu carné digital de afiliación en formato PDF. Ingresa con tu documento y contraseña, ve a la sección de certificados o carné, y descárgalo. El carné digital tiene la misma validez que el carné físico."
    },
    {
        "categoria": "CERTIFICADOS Y DOCUMENTOS",
        "pregunta": "¿Cómo generar certificados financieros o de subsidios en Comfama?",
        "respuesta": "Desde la sección de certificados en el portal de servicios de Comfama puedes generar documentos relacionados con subsidios y productos financieros. Ingresa al portal, selecciona el tipo de certificado que necesitas (financiero, subsidios, paz y salvo) y descárgalo en formato PDF."
    },
    # CUENTA DIGITAL
    {
        "categoria": "CUENTA DIGITAL",
        "pregunta": "¿Cómo crear una cuenta en Comfama?",
        "respuesta": "Para crear tu cuenta digital en Comfama, debes registrarte en el portal digital (www.comfama.com) con tu documento de identidad y correo electrónico. Sigue los pasos del formulario de registro, crea una contraseña segura y confirma tu correo electrónico. Una vez registrado, podrás acceder a todos los servicios en línea."
    },
    {
        "categoria": "CUENTA DIGITAL",
        "pregunta": "¿Cómo actualizar mi correo electrónico o celular en Comfama?",
        "respuesta": "Puedes actualizar tu correo electrónico o número de celular en la sección de perfil dentro del portal de servicios de Comfama. Ingresa con tu usuario y contraseña, ve a 'Mi perfil' o 'Datos personales' y modifica la información. También puedes hacerlo llamando a la línea de atención (604) 360 70 80."
    },
    {
        "categoria": "CUENTA DIGITAL",
        "pregunta": "Tengo problemas para iniciar sesión en Comfama, ¿qué debo hacer?",
        "respuesta": "Si tienes problemas para iniciar sesión en el portal de Comfama, utiliza la opción de 'Recuperar contraseña' en la página de inicio de sesión. Si el problema persiste, comunícate con el soporte de Comfama al (604) 360 70 80 o visita una sede para recibir asistencia personalizada."
    },
    {
        "categoria": "CUENTA DIGITAL",
        "pregunta": "¿Qué puedo hacer desde la plataforma digital de Comfama?",
        "respuesta": "Desde la plataforma digital de Comfama puedes: consultar servicios disponibles, inscribirte a cursos y programas educativos, gestionar y descargar certificados, acceder a programas de recreación y turismo, consultar información sobre créditos y subsidios, actualizar tus datos personales, y afiliar beneficiarios."
    },
    {
        "categoria": "CUENTA DIGITAL",
        "pregunta": "¿Cómo recupero mi contraseña de Comfama?",
        "respuesta": "Para recuperar tu contraseña de Comfama, ingresa al portal de servicios en línea (www.comfama.com) y haz clic en 'Olvidé mi contraseña'. Ingresa tu número de documento de identidad y recibirás un enlace de recuperación en tu correo electrónico registrado. Si no tienes acceso al correo, comunícate con soporte al (604) 360 70 80 para recibir asistencia."
    },
    # SUBSIDIOS
    {
        "categoria": "SUBSIDIOS",
        "pregunta": "¿Qué es el subsidio familiar de Comfama?",
        "respuesta": "El subsidio familiar es una prestación social pagada por Comfama a los trabajadores afiliados que ganan hasta 4 salarios mínimos legales mensuales vigentes (SMLMV) y que tienen personas a cargo (hijos, cónyuge, padres). Se paga en dinero (cuota monetaria) y en servicios (recreación, educación, salud, vivienda). El monto varía según la categoría de afiliación y el número de personas a cargo."
    },
    {
        "categoria": "SUBSIDIOS",
        "pregunta": "¿Cómo me postulo al subsidio de vivienda de Comfama?",
        "respuesta": "Para postularte al subsidio de vivienda de Comfama debes: ser afiliado activo a Comfama, no ser propietario de vivienda, tener ingresos familiares de hasta 4 SMLMV, y no haber recibido subsidio de vivienda anteriormente. La postulación se realiza a través del portal de Comfama o en las sedes cuando se abren convocatorias. Debes presentar documentos como cédula, certificado laboral, declaración de no poseer vivienda, y formulario de postulación."
    },
    {
        "categoria": "SUBSIDIOS",
        "pregunta": "¿Qué es el subsidio al desempleo de Comfama?",
        "respuesta": "El subsidio al desempleo de Comfama es un beneficio que brinda apoyo económico temporal y pago de aportes a seguridad social (salud y pensión) cuando una persona pierde su empleo. Este subsidio está diseñado para ayudar a los trabajadores cesantes mientras buscan una nueva oportunidad laboral."
    },
    {
        "categoria": "SUBSIDIOS",
        "pregunta": "¿Cuándo puedo postularme al subsidio al desempleo de Comfama?",
        "respuesta": "El proceso de postulación al subsidio al desempleo de Comfama se habilita una vez al mes en ciclos definidos por Comfama. Debes estar atento a las fechas de apertura de convocatoria que se publican en la página web de Comfama y en sus canales oficiales."
    },
    {
        "categoria": "SUBSIDIOS",
        "pregunta": "¿Existe lista de espera para el subsidio al desempleo de Comfama?",
        "respuesta": "No, no existe lista de espera para el subsidio al desempleo de Comfama. Las solicitudes se evalúan en cada ciclo mensual de postulación. Si no eres seleccionado en un ciclo, puedes volver a postularte en el siguiente ciclo disponible."
    },
    {
        "categoria": "SUBSIDIOS",
        "pregunta": "¿Cuáles son los requisitos para acceder al subsidio al desempleo de Comfama?",
        "respuesta": "Los requisitos para acceder al subsidio al desempleo de Comfama son: haber estado afiliado a una caja de compensación familiar, estar actualmente cesante (sin empleo), cumplir con los requisitos establecidos por la normativa vigente, estar inscrito en el servicio público de empleo, y no haber recibido el subsidio en los últimos 12 meses."
    },
    {
        "categoria": "SUBSIDIOS",
        "pregunta": "¿Qué documentos necesito para solicitar el subsidio al desempleo de Comfama?",
        "respuesta": "Para solicitar el subsidio al desempleo de Comfama necesitas: certificado de terminación laboral o carta de despido, inscripción vigente al servicio público de empleo, documento de identidad (cédula de ciudadanía), y formulario de solicitud diligenciado disponible en el portal de Comfama."
    },
    {
        "categoria": "SUBSIDIOS",
        "pregunta": "¿Qué beneficios incluye el subsidio al desempleo de Comfama?",
        "respuesta": "El subsidio al desempleo de Comfama incluye: apoyo económico temporal (cuota monetaria), pago de aportes a salud, pago de aportes a pensión, y acompañamiento laboral para ayudarte a encontrar un nuevo empleo a través del servicio de empleo de Comfama."
    },
    {
        "categoria": "SUBSIDIOS",
        "pregunta": "¿Puedo aplicar nuevamente al subsidio al desempleo de Comfama?",
        "respuesta": "Sí, puedes aplicar nuevamente al subsidio al desempleo de Comfama siempre que cumplas nuevamente con todos los requisitos establecidos y hayan pasado al menos 12 meses desde la última vez que recibiste el beneficio."
    },
    {
        "categoria": "SUBSIDIOS",
        "pregunta": "¿Qué pasa si consigo trabajo mientras recibo el subsidio al desempleo de Comfama?",
        "respuesta": "Si consigues trabajo mientras recibes el subsidio al desempleo de Comfama, debes informar a Comfama inmediatamente para suspender el beneficio. Esto es obligatorio y el incumplimiento puede generar sanciones o la obligación de devolver los recursos recibidos."
    },
    {
        "categoria": "SUBSIDIOS",
        "pregunta": "¿Cómo saber si mi postulación al subsidio al desempleo fue registrada correctamente?",
        "respuesta": "Puedes consultar el estado de tu postulación al subsidio al desempleo desde el portal de servicios de Comfama o a través de los canales de atención: línea telefónica (604) 360 70 80, WhatsApp, o visitando una sede de Comfama."
    },
    # CRÉDITOS Y SERVICIOS FINANCIEROS
    {
        "categoria": "CRÉDITOS Y SERVICIOS FINANCIEROS",
        "pregunta": "¿Cuáles son los requisitos para acceder al crédito de cuota monetaria de Comfama?",
        "respuesta": "Para acceder al crédito de cuota monetaria de Comfama debes: ser afiliado activo a Comfama, cumplir con los requisitos de edad establecidos, no tener mora en productos financieros con Comfama, tener capacidad de pago demostrable, y presentar la documentación requerida (cédula, certificado laboral, desprendibles de pago)."
    },
    {
        "categoria": "CRÉDITOS Y SERVICIOS FINANCIEROS",
        "pregunta": "¿Qué es el crédito de cuota monetaria de Comfama?",
        "respuesta": "El crédito de cuota monetaria de Comfama es un producto financiero que utiliza el subsidio de cuota monetaria como respaldo para facilitar el acceso a financiamiento. Es una línea de crédito exclusiva para afiliados que permite obtener recursos con tasas preferenciales, usando como garantía parcial el subsidio que recibes de Comfama."
    },
    {
        "categoria": "CRÉDITOS Y SERVICIOS FINANCIEROS",
        "pregunta": "¿Qué servicios financieros ofrece Comfama?",
        "respuesta": "Comfama ofrece diversos servicios financieros para sus afiliados: créditos de vivienda, créditos educativos, créditos de libre inversión, crédito de cuota monetaria, seguros, y programas de ahorro. Todos estos productos tienen tasas preferenciales para afiliados activos."
    },
    # EDUCACIÓN Y CURSOS
    {
        "categoria": "EDUCACIÓN Y CURSOS",
        "pregunta": "Tengo problemas con la inscripción a cursos de Comfama, ¿qué hago?",
        "respuesta": "Si tienes problemas con la inscripción a cursos de Comfama, debes comunicarte con el centro de atención al (604) 360 70 80 o revisar el estado de la inscripción en el portal de servicios. También puedes visitar la sede más cercana para recibir asistencia presencial con el proceso de inscripción."
    },
    {
        "categoria": "EDUCACIÓN Y CURSOS",
        "pregunta": "¿Cómo inscribirme a cursos de Comfama?",
        "respuesta": "Para inscribirte a cursos de Comfama, ingresa a la plataforma digital de Comfama (www.comfama.com), busca el curso de tu interés en la sección de educación, selecciona el curso y realiza el pago correspondiente si aplica. Algunos cursos son gratuitos para afiliados. También puedes inscribirte presencialmente en las sedes de Comfama."
    },
    {
        "categoria": "EDUCACIÓN Y CURSOS",
        "pregunta": "¿Qué programas educativos ofrece Comfama?",
        "respuesta": "Comfama ofrece una amplia variedad de programas educativos: cursos técnicos y tecnológicos, formación laboral y para el empleo, programas culturales y artísticos, educación continua y diplomados, talleres y seminarios, programas de emprendimiento, y convenios con instituciones educativas para acceder a descuentos en matrículas."
    },
    {
        "categoria": "EDUCACIÓN Y CURSOS",
        "pregunta": "¿Los cursos de Comfama tienen costo?",
        "respuesta": "Comfama ofrece tanto cursos gratuitos como cursos con costo. Muchos programas de formación básica, talleres culturales y actividades recreativas son gratuitos para afiliados. Los cursos técnicos, diplomados y programas especializados pueden tener un costo, pero los afiliados a Comfama reciben descuentos significativos. Puedes consultar la oferta educativa y los precios en www.comfama.com o llamando al (604) 360 70 80."
    },
    # ATENCIÓN Y SOPORTE
    {
        "categoria": "ATENCIÓN Y SOPORTE",
        "pregunta": "¿Cómo puedo comunicarme con Comfama?",
        "respuesta": "Puedes comunicarte con Comfama a través de múltiples canales: teléfono regional (604) 360 70 80, línea nacional gratuita 018000 415 455, WhatsApp, chat digital en la página web www.comfama.com, o visitando cualquiera de las sedes físicas de Comfama en Medellín y municipios de Antioquia."
    },
    {
        "categoria": "ATENCIÓN Y SOPORTE",
        "pregunta": "¿Cuál es el número de atención de Comfama?",
        "respuesta": "Los números de atención de Comfama son: línea regional (604) 360 70 80 para llamadas desde Medellín y el área metropolitana, y línea nacional gratuita 018000 415 455 para llamadas desde cualquier parte de Colombia. El horario de atención es de lunes a viernes de 7:00 a.m. a 7:00 p.m. y sábados de 8:00 a.m. a 1:00 p.m."
    },
    {
        "categoria": "ATENCIÓN Y SOPORTE",
        "pregunta": "¿Cómo presentar una petición, queja o reclamo en Comfama?",
        "respuesta": "Para presentar una petición, queja o reclamo (PQR) en Comfama, debes utilizar el formulario de contacto disponible en el sitio web www.comfama.com en la sección de 'Contáctanos' o 'PQR'. También puedes hacerlo presencialmente en cualquier sede, por teléfono al (604) 360 70 80, o por correo electrónico."
    },
    {
        "categoria": "ATENCIÓN Y SOPORTE",
        "pregunta": "¿Cuál es el correo electrónico de Comfama?",
        "respuesta": "Comfama tiene correos institucionales habilitados para casos específicos como línea de transparencia y notificaciones judiciales. Para consultas generales, se recomienda usar el formulario de contacto en www.comfama.com, llamar al (604) 360 70 80, o usar el chat digital disponible en la página web."
    },
    {
        "categoria": "ATENCIÓN Y SOPORTE",
        "pregunta": "¿Qué pasos debo seguir al llamar a la línea de atención de Comfama?",
        "respuesta": "Al llamar a la línea de atención de Comfama al (604) 360 70 80 o 018000 415 455, debes seleccionar la opción correspondiente según el servicio que necesites en el menú interactivo. Ten a la mano tu número de documento de identidad para agilizar la atención. Un asesor te guiará según tu consulta."
    },
    {
        "categoria": "ATENCIÓN Y SOPORTE",
        "pregunta": "¿Cuáles son los horarios de atención de Comfama?",
        "respuesta": "Los horarios de atención de Comfama son: línea telefónica (604) 360 70 80 de lunes a viernes de 7:00 a.m. a 7:00 p.m. y sábados de 8:00 a.m. a 1:00 p.m. Las sedes físicas atienden de lunes a viernes de 7:30 a.m. a 5:00 p.m. y sábados de 8:00 a.m. a 12:00 m. Los horarios pueden variar según la sede. El chat digital y WhatsApp están disponibles en horario extendido."
    },
]

# Escribir cada FAQ como sección independiente
current_category = ""
for faq in faqs:
    # Agregar encabezado de categoría si cambió
    if faq["categoria"] != current_category:
        current_category = faq["categoria"]
        doc.add_heading(current_category, level=1)
        doc.add_paragraph("")  # Separador
    
    # Pregunta en negrita
    p_question = doc.add_paragraph()
    run = p_question.add_run(f"Pregunta: {faq['pregunta']}")
    run.bold = True
    run.font.size = Pt(12)
    
    # Respuesta
    p_answer = doc.add_paragraph()
    run = p_answer.add_run(f"Respuesta: {faq['respuesta']}")
    run.font.size = Pt(11)
    
    # Separador claro entre FAQs
    doc.add_paragraph("---")

# Guardar
output_path = "knowledge-base-docs/FAQ_Comfama_Centro_Conocimiento_v2.docx"
doc.save(output_path)
print(f"✅ Documento generado: {output_path}")
print(f"📊 Total FAQs: {len(faqs)}")
